from docxtpl import DocxTemplate
import os
from typing import Dict, Any
from docx import Document

class LessonRenderer:
    def __init__(self, template_path: str = "templates/lesson_plan_template.docx"):
        self.template_path = template_path
    def render(self, data: Dict[str, Any], output_path: str):
        """
        data 结构兼容性处理：支持直接 content 或结构化的 sections/steps
        """
        # --- 情况 A: 存在预设模板 (使用 docxtpl) ---
        if os.path.exists(self.template_path):
            try:
                doc = DocxTemplate(self.template_path)
                doc.render(data)
                doc.save(output_path)
                return
            except Exception as e:
                print(f"⚠️ 模板渲染报错，切换至基础生成: {e}")

        # --- 情况 B: 无模板或渲染失败 (使用 python-docx 动态生成) ---
        doc = Document()
        
        # 1. 写入标题
        doc.add_heading(data.get('title', '教学设计方案'), 0)

        # 2. 核心逻辑：遍历内容
        # 优先查找结构化的 sections (来自最新的解析逻辑)
        sections = data.get('sections', [])
        steps = data.get('steps', [])
        
        if sections:
            for sec in sections:
                doc.add_heading(sec.get('heading', ''), level=1)
                doc.add_paragraph(sec.get('content', ''))
        
        # 兼容旧的表格逻辑 (steps)
        elif steps:
            for step in steps:
                phase = step.get('phase', '')
                act = step.get('teacher_act', '')
                doc.add_heading(phase, level=2)
                doc.add_paragraph(act)

        # 3. 兜底逻辑：如果都没有，直接找长文本字段
        else:
            # 依次尝试获取可能存储全文的键
            raw_text = data.get('content') or data.get('markdown') or data.get('dsl')
            if raw_text:
                doc.add_paragraph(str(raw_text))
            else:
                doc.add_paragraph("未匹配到详细教案内容，请检查 Dify 输出。")

        # 4. 保存
        doc.save(output_path)